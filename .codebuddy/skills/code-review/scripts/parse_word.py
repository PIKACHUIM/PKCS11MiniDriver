#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word 文档需求解析脚本

解析 .docx 格式的需求文档
"""

import sys
import argparse
import subprocess
from pathlib import Path

script_dir = Path(__file__).parent
sys.path.insert(0, str(script_dir))


def ensure_python_docx() -> bool:
    """确保 python-docx 可用；如果未安装则自动安装。

    Returns:
        是否已可用
    """
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        pass

    # 优先用当前解释器安装，避免环境不一致
    candidates = [
        [sys.executable, "-m", "pip", "install", "python-docx"],
        ["pip", "install", "python-docx"],
        ["pip3", "install", "python-docx"],
    ]

    for cmd in candidates:
        try:
            subprocess.check_call(cmd)
            import docx  # noqa: F401
            return True
        except Exception:
            continue

    return False


try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    # 尝试自动安装并重试导入
    HAS_DOCX = ensure_python_docx()
    if HAS_DOCX:
        from docx import Document


def parse_word_content(file_path: str) -> str:
    """解析 Word 文档并返回全文内容字符串
    
    Args:
        file_path: Word 文档路径
    
    Returns:
        解析后的全文内容字符串（段落与表格行以换行拼接）
    """
    if not HAS_DOCX:
        print("警告: python-docx 未安装且自动安装失败，无法解析 Word 文档", file=sys.stderr)
        print("请手动运行: pip install python-docx", file=sys.stderr)
        return ""
    
    if not Path(file_path).exists():
        print(f"错误: 文件不存在: {file_path}", file=sys.stderr)
        return ""
    
    try:
        doc = Document(file_path)
        
        # 提取全文内容
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    full_text.append(row_text)
        
        return '\n'.join(full_text)
        
    except Exception as e:
        print(f"错误: 解析 Word 文档失败: {e}", file=sys.stderr)
        return ""


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='解析 Word 需求文档')
    parser.add_argument('file', help='Word 文档路径')
    args = parser.parse_args()
    
    content = parse_word_content(args.file)
    if not content:
        return 1

    print(content)
    return 0


if __name__ == '__main__':
    sys.exit(main())